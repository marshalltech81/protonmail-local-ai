"""Tests for src/extractors/.

Covers the dispatch (MIME → module → extension fallback → unsupported),
the simple text + html extractors end-to-end, the PDF digital path
against a tiny synthesized PDF, and the safety properties the
dispatcher itself enforces (size cap, OCR gate, exception → ``failed``).

OCR-dependent paths (Tesseract, Poppler) are tested via mocks rather
than against a live binary so the test suite stays runnable on a
laptop without the Docker image's apt packages installed.
"""

from __future__ import annotations

import pytest
from src.extractors import (
    STATUS_EMPTY,
    STATUS_FAILED,
    STATUS_SUCCESS,
    STATUS_TOO_LARGE,
    STATUS_UNSUPPORTED,
    ExtractionResult,
    extract,
)


class TestDispatchByMime:
    def test_text_plain_routes_to_text_extractor(self):
        result = extract(
            content_type="text/plain",
            filename="note.txt",
            payload=b"Hello there",
        )
        assert result.status == STATUS_SUCCESS
        assert result.extractor == "text"
        assert result.text == "Hello there"

    def test_text_csv_uses_text_extractor(self):
        result = extract(
            content_type="text/csv",
            filename="data.csv",
            payload=b"a,b,c\n1,2,3",
        )
        assert result.status == STATUS_SUCCESS
        assert "1,2,3" in result.text

    def test_text_html_renders_via_html_extractor(self):
        result = extract(
            content_type="text/html",
            filename="page.html",
            payload=b"<html><body><h1>Title</h1><p>Body text.</p></body></html>",
        )
        assert result.status == STATUS_SUCCESS
        assert result.extractor == "html"
        assert "Title" in result.text
        assert "Body text" in result.text

    def test_unknown_mime_with_known_extension_falls_back(self):
        # ``application/octet-stream`` is the catch-all clients use when
        # MIME detection fails. Filename extension routing must rescue
        # these.
        result = extract(
            content_type="application/octet-stream",
            filename="reading.txt",
            payload=b"text content",
        )
        assert result.status == STATUS_SUCCESS
        assert result.extractor == "text"

    def test_no_dispatch_match_returns_unsupported(self):
        result = extract(
            content_type="application/x-unknown",
            filename="mystery.bin",
            payload=b"\x00\x01",
        )
        assert result.status == STATUS_UNSUPPORTED
        assert result.text is None


class TestSafetyGates:
    def test_payload_over_max_bytes_returns_too_large(self):
        result = extract(
            content_type="text/plain",
            filename="huge.txt",
            payload=b"x" * 200,
            max_bytes=100,
        )
        assert result.status == STATUS_TOO_LARGE
        assert result.text is None
        assert "200" in (result.error or "")

    def test_image_dispatch_blocked_when_ocr_disabled(self):
        result = extract(
            content_type="image/png",
            filename="screenshot.png",
            payload=b"\x89PNG\r\n\x1a\n",
            ocr_enabled=False,
        )
        # Without OCR the image extractor has nothing useful to do —
        # downgrade to unsupported so a future OCR-enabled re-run can
        # upgrade the cached row.
        assert result.status == STATUS_UNSUPPORTED
        assert "OCR disabled" in (result.error or "")

    def test_extractor_exception_becomes_failed_not_raised(self, monkeypatch):
        # Dispatcher must convert per-format exceptions into a ``failed``
        # ExtractionResult so a single malformed attachment cannot
        # dead-letter the parent message's indexing job.
        from src.extractors import _safe_import as real_safe_import

        def fake_safe_import(module_name):
            if module_name == "text":

                def boom(payload, **opts):
                    raise RuntimeError("simulated extractor crash")

                return boom
            return real_safe_import(module_name)

        monkeypatch.setattr("src.extractors._safe_import", fake_safe_import)
        # Bust the import cache so the fake gets installed.
        monkeypatch.setattr("src.extractors._IMPORT_CACHE", {})

        result = extract(
            content_type="text/plain",
            filename="boom.txt",
            payload=b"hi",
        )
        assert result.status == STATUS_FAILED
        assert "simulated extractor crash" in result.error
        assert result.extractor == "text"


class TestEmptyExtraction:
    def test_extractor_returning_empty_text_reports_empty(self):
        # ``text`` extractor on whitespace-only payload yields a
        # whitespace-stripped empty string → status="empty", not
        # "success" with a blank text field.
        result = extract(
            content_type="text/plain",
            filename="blank.txt",
            payload=b"   \n\t  ",
        )
        assert result.status == STATUS_EMPTY
        assert result.text is None


class TestExtractionResultShape:
    def test_dataclass_is_frozen(self):
        r = ExtractionResult(status=STATUS_SUCCESS, extractor="x", text="y", error=None)
        with pytest.raises(Exception):
            r.status = STATUS_EMPTY  # type: ignore[misc]


# ---------------------------------------------------------------------------
# Per-format extractors — small synthetic fixtures rather than file
# fixtures so the tests stay readable and the test suite stays self-
# contained. OCR-dependent paths use mocks for ``pytesseract`` and
# ``pdf2image`` so the tests run without the Docker image's apt
# packages installed on the host.
# ---------------------------------------------------------------------------


class TestTextExtractorFallback:
    def test_invalid_utf8_falls_back_to_replace_decode(self):
        """``text`` extractor must not raise on ill-formed bytes — the
        chunker can still index a payload with replacement characters,
        but a hard decode failure would dead-letter the parent message.
        """
        from src.extractors.text import extract as text_extract

        # ``\xff\xfe`` is an invalid UTF-8 start byte sequence in
        # context. The replacement-decode fallback yields valid Unicode.
        payload = b"before" + b"\xff\xfe" + b"after"
        text, name = text_extract(payload)
        assert name == "text"
        assert "before" in text
        assert "after" in text


class TestHtmlExtractorFallback:
    def test_invalid_utf8_html_falls_back_to_replace_decode(self):
        from src.extractors.html import extract as html_extract

        payload = b"<html><body>" + b"\xff\xfe" + b"text</body></html>"
        text, name = html_extract(payload)
        assert name == "html"
        assert "text" in text


class TestDocxExtractor:
    def test_extracts_paragraphs_and_tables(self):
        import io

        import docx
        from src.extractors.docx import extract as docx_extract

        document = docx.Document()
        document.add_paragraph("Invoice number 12345")
        document.add_paragraph("Due date: 2024-04-30")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Vendor"
        table.cell(0, 1).text = "Amount"
        table.cell(1, 0).text = "Acme Corp"
        table.cell(1, 1).text = "$500"
        buf = io.BytesIO()
        document.save(buf)

        text, name = docx_extract(buf.getvalue())
        assert name == "docx"
        assert "Invoice number 12345" in text
        assert "Due date: 2024-04-30" in text
        assert "Vendor Amount" in text
        assert "Acme Corp $500" in text


class TestXlsxExtractor:
    def test_serializes_each_sheet_with_header_marker(self):
        import io

        import openpyxl
        from src.extractors.xlsx import extract as xlsx_extract

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Q1"
        ws.append(["Item", "Price"])
        ws.append(["Widget", 25])
        ws.append(["Gadget", 75])
        buf = io.BytesIO()
        wb.save(buf)
        wb.close()

        text, name = xlsx_extract(buf.getvalue())
        assert name == "xlsx"
        assert "[Sheet: Q1]" in text
        assert "Item\tPrice" in text
        assert "Widget\t25" in text


class TestPdfDigitalExtractor:
    def test_extracts_text_from_minimal_digital_pdf(self):
        """A small synthetic PDF with a real text layer must round-trip
        through the digital path without invoking OCR. ``pypdf`` itself
        is the canonical PDF builder available — synthesizing a valid
        PDF byte stream by hand is too brittle, so we use pypdf to
        write and pypdf to read.
        """

        from pathlib import Path as _P

        from src.extractors.pdf import extract as pdf_extract

        fixture = _P(__file__).parent / "fixtures" / "extractors" / "digital.pdf"
        if not fixture.exists():
            # Generate the fixture on demand the first time the test
            # runs. Subsequent runs reuse it for determinism.
            fixture.parent.mkdir(parents=True, exist_ok=True)
            from pypdf import PdfWriter
            from pypdf.generic import (
                ContentStream,
                DictionaryObject,
                NameObject,
            )

            writer = PdfWriter()
            page = writer.add_blank_page(width=612, height=792)
            # Minimal text-layer content stream: BT / Tf / Td / Tj / ET.
            stream = ContentStream(None, writer)
            stream._data = b"BT /F1 12 Tf 72 720 Td (Invoice number 42) Tj ET"
            page[NameObject("/Contents")] = stream
            # Register a basic Type 1 font so the Tf reference resolves.
            font = DictionaryObject(
                {
                    NameObject("/Type"): NameObject("/Font"),
                    NameObject("/Subtype"): NameObject("/Type1"),
                    NameObject("/BaseFont"): NameObject("/Helvetica"),
                }
            )
            resources = DictionaryObject(
                {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
            )
            page[NameObject("/Resources")] = resources
            with fixture.open("wb") as f:
                writer.write(f)

        payload = fixture.read_bytes()
        text, name = pdf_extract(payload, ocr_enabled=False)
        # The fixture's content stream contains a literal "Invoice
        # number 42". Assert the digital path round-trips it so a
        # regression in the pypdf pin or in ``_extract_digital``
        # surfaces here rather than silently degrading retrieval.
        assert "Invoice number 42" in text
        assert name == "pdf-digital"

    def test_ocr_fallback_invoked_when_digital_too_short(self, monkeypatch):
        """When ``_extract_digital`` returns near-empty text (a scanned
        PDF), the dispatcher must call into the OCR path. Mocked here to
        avoid requiring Tesseract + Poppler at test time.
        """
        from src.extractors import pdf

        monkeypatch.setattr(pdf, "_extract_digital", lambda payload: "")
        monkeypatch.setattr(
            pdf,
            "_extract_ocr",
            lambda payload, *, max_ocr_pages: "OCR'd page 1\n\nOCR'd page 2",
        )

        text, name = pdf.extract(b"%PDF-1.7 (mocked)", ocr_enabled=True, max_ocr_pages=5)
        assert name == "pdf-ocr"
        assert "OCR'd page 1" in text

    def test_ocr_fallback_failure_is_reported_as_failed(self, monkeypatch):
        """If a scanned PDF needs OCR but Poppler/Tesseract fails, the
        dispatcher should cache a failed extraction with the error
        instead of treating near-empty digital text as a successful
        no-text result.
        """
        from src.extractors import pdf

        monkeypatch.setattr(pdf, "_extract_digital", lambda payload: "")

        def fail_ocr(payload, *, max_ocr_pages):
            raise RuntimeError("poppler missing")

        monkeypatch.setattr(pdf, "_extract_ocr", fail_ocr)

        result = extract(
            content_type="application/pdf",
            filename="scan.pdf",
            payload=b"%PDF-1.7",
        )

        assert result.status == STATUS_FAILED
        assert result.extractor == "pdf"
        assert result.text is None
        assert "poppler missing" in (result.error or "")

    def test_ocr_disabled_returns_digital_text_only(self, monkeypatch):
        from src.extractors import pdf

        monkeypatch.setattr(pdf, "_extract_digital", lambda payload: "tiny")
        # OCR must NOT be called when ocr_enabled=False, even if digital
        # text is too short to satisfy ``_MIN_DIGITAL_CHARS``.
        ocr_called = []
        monkeypatch.setattr(
            pdf,
            "_extract_ocr",
            lambda *a, **kw: ocr_called.append(True) or "",
        )

        text, name = pdf.extract(b"%PDF-1.7", ocr_enabled=False, max_ocr_pages=5)
        assert text == "tiny"
        assert name == "pdf-digital"
        assert ocr_called == []


class TestImageExtractor:
    def test_invokes_pytesseract_with_oriented_image(self, monkeypatch):
        import io

        from PIL import Image
        from src.extractors import image as image_module

        captured = {}

        def fake_image_to_string(img):
            captured["called"] = True
            captured["mode"] = img.mode
            return "RECEIPT TOTAL $42.00\n"

        monkeypatch.setattr(image_module.pytesseract, "image_to_string", fake_image_to_string)

        # Build a tiny in-memory image so the extractor has real bytes
        # to load through PIL.
        buf = io.BytesIO()
        Image.new("RGB", (10, 10), color="white").save(buf, format="PNG")
        text, name = image_module.extract(buf.getvalue())

        assert captured["called"] is True
        assert "RECEIPT TOTAL" in text
        assert name == "image-ocr"

    def test_exif_orientation_rotates_image_before_ocr(self, monkeypatch):
        """An image with EXIF Orientation=6 (rotated 90° CW for display)
        must reach pytesseract post-rotation. ``ImageOps.exif_transpose``
        is the production mechanism — verify it actually fires by
        building a wide source image and asserting the OCR'd image
        comes through with the rotated dimensions."""
        import io

        from PIL import Image
        from src.extractors import image as image_module

        # Wide source image (40x10). Orientation=6 means "rotate 90° CW
        # for display", so post-transpose the image becomes 10×40.
        src = Image.new("RGB", (40, 10), color="white")
        exif = src.getexif()
        # 0x0112 is the standard Orientation tag. Pillow's
        # ``Image.Exif`` accepts integer keys directly, avoiding a new
        # piexif dependency just for the test.
        exif[0x0112] = 6
        buf = io.BytesIO()
        src.save(buf, format="JPEG", exif=exif.tobytes())

        captured: dict[str, tuple[int, int]] = {}

        def fake_image_to_string(img):
            captured["size"] = img.size
            return "ok"

        monkeypatch.setattr(image_module.pytesseract, "image_to_string", fake_image_to_string)

        image_module.extract(buf.getvalue())

        # Post-rotation the image is 10 wide × 40 tall — assert we did
        # not OCR the unrotated 40×10 source.
        assert captured["size"] == (10, 40)

    def test_decompression_bomb_warning_promoted_to_error(self, monkeypatch):
        """A canvas in the warning band (1×–2× the cap) must surface
        as a raised ``DecompressionBombWarning`` (promoted to error
        inside the ``warnings.catch_warnings()`` scope) so the
        dispatcher records it as ``failed`` rather than OOM'ing the
        worker. PIL raises ``DecompressionBombError`` past 2× directly,
        so we size the test image into the warning band only."""
        import io

        import pytest
        from PIL import Image
        from src.extractors import image as image_module

        # 50×50 = 2500 pixels. Cap at 1500 puts the image at 1.67× the
        # cap — inside the warning band (Error fires only past 2×).
        monkeypatch.setattr(image_module, "_MAX_IMAGE_PIXELS", 1500)

        buf = io.BytesIO()
        Image.new("RGB", (50, 50), color="white").save(buf, format="PNG")

        with pytest.raises(Image.DecompressionBombWarning):
            image_module.extract(buf.getvalue())

    def test_decompression_bomb_error_propagates(self, monkeypatch):
        """A canvas past 2× the cap must surface PIL's
        ``DecompressionBombError`` directly. The dispatcher's exception
        handler converts both this and the warning-band raise into a
        ``failed`` ExtractionResult."""
        import io

        import pytest
        from PIL import Image
        from src.extractors import image as image_module

        # 50×50 = 2500 pixels. Cap at 100 → 25× the cap → Error.
        monkeypatch.setattr(image_module, "_MAX_IMAGE_PIXELS", 100)

        buf = io.BytesIO()
        Image.new("RGB", (50, 50), color="white").save(buf, format="PNG")

        with pytest.raises(Image.DecompressionBombError):
            image_module.extract(buf.getvalue())

    def test_max_image_pixels_restored_after_extract(self, monkeypatch):
        """``Image.MAX_IMAGE_PIXELS`` must be restored after a call so
        other PIL consumers in the same process aren't permanently
        constrained by the extractor's cap."""
        import io

        from PIL import Image
        from src.extractors import image as image_module

        prior = Image.MAX_IMAGE_PIXELS
        monkeypatch.setattr(image_module.pytesseract, "image_to_string", lambda img: "ok")

        buf = io.BytesIO()
        Image.new("RGB", (10, 10), color="white").save(buf, format="PNG")
        image_module.extract(buf.getvalue())

        assert Image.MAX_IMAGE_PIXELS == prior


class TestDispatcherTextCap:
    def test_max_extracted_chars_truncates_success_text(self):
        result = extract(
            content_type="text/plain",
            filename="long.txt",
            payload=b"abcdefghijklmnopqrstuvwxyz" * 100,
            max_extracted_chars=64,
        )
        assert result.status == STATUS_SUCCESS
        assert result.text is not None
        assert len(result.text) == 64

    def test_max_extracted_chars_none_does_not_truncate(self):
        payload = b"abcdefghij" * 50
        result = extract(
            content_type="text/plain",
            filename="long.txt",
            payload=payload,
            max_extracted_chars=None,
        )
        assert result.status == STATUS_SUCCESS
        assert result.text is not None
        assert len(result.text) == len(payload)
