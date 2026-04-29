"""DOCX (Word .docx) extractor.

Uses ``python-docx`` to walk paragraphs and tables. Tables are
serialized cell-by-cell separated by single spaces so a row's cells
read together for retrieval, while preserving paragraph structure
elsewhere so the downstream chunker has paragraph boundaries to pack
on. Header / footer text is included because invoices and contracts
routinely place key fields (vendor name, dates, totals) there.

Legacy ``.doc`` (binary Word, not OOXML) cannot be parsed by
``python-docx``; the dispatcher routes those to this module too but
the call will raise ``BadZipFile`` and surface as ``failed`` —
acceptable until / unless a real ``.doc`` extractor (e.g. ``antiword``,
``catdoc``) is added.
"""

from __future__ import annotations

import io

import docx as _docx


def extract(
    payload: bytes,
    *,
    ocr_enabled: bool = True,  # noqa: ARG001
    max_ocr_pages: int = 20,  # noqa: ARG001
    ocr_timeout_seconds: float | None = None,  # noqa: ARG001
    max_pdf_pages: int | None = None,  # noqa: ARG001
) -> tuple[str, str]:
    """Extract text from a DOCX payload. Returns (text, "docx")."""
    document = _docx.Document(io.BytesIO(payload))

    parts: list[str] = []

    # Body paragraphs preserve the document's natural paragraph
    # structure — the chunker keys off blank-line gaps between paragraphs.
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Tables: serialize each row as space-joined cells so a header row
    # like "Invoice #  Date  Amount" stays on one line and matches a
    # search for any of those tokens. Empty cells are dropped from the
    # row to avoid runs of double-spaces that would dilute FTS scoring.
    for table in document.tables:
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells]
            cell_texts = [c for c in cell_texts if c]
            if cell_texts:
                parts.append(" ".join(cell_texts))

    # Headers + footers (per section). Many real-world invoices put
    # vendor name and address in the header, totals in the footer; if
    # we ignore them, the most retrieval-relevant fields go missing.
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for paragraph in section.footer.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

    return "\n\n".join(parts), "docx"
